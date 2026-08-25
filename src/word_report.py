from __future__ import annotations

from io import BytesIO
from datetime import datetime
from typing import Callable

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "18324A"
BLUE = "2E75B6"
LIGHT_BLUE = "DCEAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
BLACK = "111827"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def _set_font(run, size=9.5, bold=False, color=BLACK):
    run.font.name = "Hiragino Sans GB"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_table(doc, headers, rows, widths, numeric_columns=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(str(header)), size=8.5, bold=True, color=WHITE)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            if row_index % 2:
                _shade(cells[col_index], LIGHT_GRAY)
            p = cells[col_index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT
            _set_font(p.add_run(str(value)), size=8.2)
    _set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def _add_hyperlink(paragraph, text, url):
    if not url:
        return
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _fmt(value, digits=2, suffix=""):
    try:
        if pd.isna(value):
            return "—"
        return f"{float(value):.{digits}f}{suffix}"
    except (TypeError, ValueError):
        return "—"


def _text(value, fallback="—"):
    return fallback if pd.isna(value) or not str(value).strip() else str(value)


def _chart(values, title, value_label="比重 %"):
    """Create a compact Word-ready chart; numbered labels avoid CJK font loss on Linux."""
    from PIL import Image, ImageDraw, ImageFont

    data = [(str(name), float(value)) for name, value in values if pd.notna(value)]
    if not data:
        return None
    numbers = [item[1] for item in data]
    width, height = 1125, 470
    left, right, top, bottom = 70, 30, 70, 65
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()
    draw.text((left, 18), title, fill="#18324A", font=font)
    draw.text((8, top), value_label, fill="#667085", font=font)
    plot_width, plot_height = width - left - right, height - top - bottom
    max_value = max(numbers) or 1
    for grid_index in range(5):
        y = top + plot_height * grid_index / 4
        draw.line((left, y, width - right, y), fill="#D7DEE7", width=1)
        grid_value = max_value * (1 - grid_index / 4)
        draw.text((18, y - 6), f"{grid_value:.1f}", fill="#667085", font=font)
    slot = plot_width / len(numbers)
    bar_width = max(12, slot * 0.58)
    for idx, number in enumerate(numbers, start=1):
        x0 = left + (idx - 1) * slot + (slot - bar_width) / 2
        x1 = x0 + bar_width
        y0 = top + plot_height * (1 - number / max_value)
        draw.rectangle((x0, y0, x1, top + plot_height), fill="#7CB9E8", outline="#2E75B6")
        draw.text((x0 + 2, max(top, y0 - 15)), f"{number:.1f}", fill="#18324A", font=font)
        draw.text((x0 + bar_width / 2 - 7, top + plot_height + 12), f"#{idx}", fill="#111827", font=font)
    output = BytesIO()
    image.save(output, format="PNG", optimize=True)
    output.seek(0)
    return output


def _setup_document(category, updated_at, sort_label):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 12, 6),
        ("Heading 2", 12.5, BLUE, 9, 4),
        ("Heading 3", 10.5, NAVY, 7, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(header.add_run(f"全球共同基金篩選器｜{category}"), size=8, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(footer.add_run("內部教育訓練使用｜資料僅供參考，不構成投資建議"), size=7.5, color=MID_GRAY)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    _set_font(kicker.add_run("MARKET / THEME FUND REPORT"), size=9, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    _set_font(title.add_run(f"{category}基金研究報告"), size=25, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_font(subtitle.add_run("基金績效、動能趨勢、產業配置與主要持股"), size=12, color=MID_GRAY)
    _add_table(
        doc,
        ["報告日期", "資料更新", "排序方式", "篩選口徑"],
        [[datetime.now().strftime("%Y-%m-%d"), updated_at or "—", sort_label, "沿用網頁左側目前條件"]],
        [1900, 2500, 2460, 2500],
    )
    note = doc.add_paragraph()
    _set_font(note.add_run("用途說明｜"), size=8.5, bold=True, color=NAVY)
    _set_font(note.add_run("本報告為內部教育訓練資料；績效為歷史資料，配息可能來自收益平準金，實際資訊以基金公司最新公告為準。"), size=8.5, color=MID_GRAY)
    return doc


def _add_performance_section(doc, funds, registration_type: Callable):
    doc.add_heading("一、基金績效與動能趨勢", level=1)
    p = doc.add_paragraph("動能加速度 = 3M 年化動能 − 6M 年化動能；正值代表近期動能相對加速，負值代表漲勢放緩或轉弱。")
    p.style = doc.styles["Normal"]
    rows = []
    for _, fund in funds.head(10).iterrows():
        acceleration = fund.get("momentum_acceleration")
        trend = "↑" if pd.notna(acceleration) and acceleration > 0 else "↓" if pd.notna(acceleration) and acceleration < 0 else "→"
        rows.append([
            fund.get("name", "—"), registration_type(fund.get("moneydj_id")), fund.get("nav_change_display", "—"),
            _fmt(fund.get("return_1m")), _fmt(fund.get("return_3m")), _fmt(fund.get("momentum_6m")),
            _fmt(fund.get("return_1y")), f"{trend} {_fmt(acceleration)}", _fmt(fund.get("sharpe"), 3), str(fund.get("data_date") or "—"),
        ])
    _add_table(
        doc,
        ["基金", "身分", "日變動", "1M%", "3M%", "6M%", "1Y%", "加速度", "夏普", "日期"],
        rows,
        [2350, 850, 850, 650, 650, 650, 650, 800, 650, 860],
        numeric_columns=(3, 4, 5, 6, 7, 8),
    )
    risk_rows = []
    for _, fund in funds.head(10).iterrows():
        risk_rows.append([
            fund.get("name", "—"), _text(fund.get("benchmark")),
            _fmt(fund.get("benchmark_return_1y")), _fmt(fund.get("excess_return_1y")),
            _fmt(fund.get("max_drawdown")), _text(fund.get("signal")),
        ])
    _add_table(
        doc,
        ["基金", "Benchmark", "Benchmark 1Y%", "超額報酬%", "最大回撤%", "訊號"],
        risk_rows,
        [2700, 2200, 1250, 1150, 1150, 910],
        numeric_columns=(2, 3, 4),
    )
    for _, fund in funds.head(10).iterrows():
        url = fund.get("tcb_url") or ""
        if url:
            link_p = doc.add_paragraph()
            link_p.paragraph_format.space_after = Pt(2)
            _set_font(link_p.add_run(f"{fund.get('name')}："), size=8.2)
            _add_hyperlink(link_p, "開啟基金績效頁", url)


def build_category_report(
    category: str,
    funds: pd.DataFrame,
    portfolio: pd.DataFrame,
    updated_at: str,
    sort_label: str,
    country_categories: set[str],
    registration_type: Callable,
    stock_url: Callable,
    direct_products: list[dict] | None = None,
) -> bytes:
    doc = _setup_document(category, updated_at, sort_label)
    _add_performance_section(doc, funds, registration_type)

    if category in country_categories:
        doc.add_heading("二、閱讀重點", level=1)
        doc.add_paragraph("國家與區域類別依網頁規則呈現基金績效、Benchmark 與風險指標，不列個股持股。請同時留意匯率、區域政策及市場估值變化。")
    else:
        doc.add_heading("二、基金 Top 10 與相關持股摘要", level=1)
        summary_rows = []
        for _, fund in funds.head(10).iterrows():
            fund_portfolio = portfolio[portfolio["fund"].eq(fund.get("name"))] if not portfolio.empty else pd.DataFrame()
            holdings = fund_portfolio[fund_portfolio["kind"].eq("holding")].sort_values("weight", ascending=False) if not fund_portfolio.empty else pd.DataFrame()
            dates = holdings["data_date"].dropna().astype(str) if not holdings.empty else pd.Series(dtype=str)
            summary_rows.append([
                fund.get("name"), registration_type(fund.get("moneydj_id")),
                fund.get("distribution") if pd.notna(fund.get("distribution")) and str(fund.get("distribution")).strip() else "待確認",
                fund.get("nav_change_display") or "—", _fmt(fund.get("return_1y")),
                _fmt(holdings["weight"].sum() if not holdings.empty else np.nan),
                "、".join(holdings["name"].astype(str).head(3)) if not holdings.empty else "待抓取",
                dates.max() if not dates.empty else "—",
            ])
        _add_table(doc, ["基金", "身分", "配息", "日變動", "1Y%", "揭露持股%", "主要持股", "持股日期"], summary_rows,
                   [2200, 850, 700, 850, 600, 800, 2360, 1000], numeric_columns=(4, 5))

        if direct_products:
            doc.add_heading("三、直接商品價格掛鉤產品", level=1)
            rows = [[item.get("產品"), item.get("代號"), item.get("直接掛鉤"), item.get("主要持有資產"), item.get("產品類型")] for item in direct_products]
            _add_table(doc, ["產品", "代號", "直接掛鉤", "主要持有資產", "類型"], rows, [1900, 650, 1800, 3510, 1500])
            for item in direct_products:
                p = doc.add_paragraph()
                _set_font(p.add_run(f"{item.get('產品')}："), size=8.2)
                _add_hyperlink(p, "官方資料", item.get("官方資料", ""))

        section_number = 4 if direct_products else 3
        doc.add_heading(f"{section_number_label(section_number)}、各基金產業配置與前十大持股", level=1)
        category_portfolio = portfolio[portfolio["fund"].isin(funds["name"])] if not portfolio.empty else pd.DataFrame()
        for fund_index, (_, fund) in enumerate(funds.head(10).iterrows(), start=1):
            if fund_index > 1:
                doc.add_page_break()
            name = str(fund.get("name"))
            fund_portfolio = category_portfolio[category_portfolio["fund"].eq(name)] if not category_portfolio.empty else pd.DataFrame()
            industries = fund_portfolio[fund_portfolio["kind"].eq("industry")].sort_values("weight", ascending=False).head(10) if not fund_portfolio.empty else pd.DataFrame()
            holdings = fund_portfolio[fund_portfolio["kind"].eq("holding")].sort_values("weight", ascending=False).head(10) if not fund_portfolio.empty else pd.DataFrame()
            dates = fund_portfolio["data_date"].dropna().astype(str) if not fund_portfolio.empty else pd.Series(dtype=str)
            doc.add_heading(f"{fund_index}. {name}", level=2)
            meta = doc.add_paragraph()
            _set_font(meta.add_run(f"資料日期：{dates.max() if not dates.empty else '—'}｜"), size=8.5, color=MID_GRAY)
            _add_hyperlink(meta, "基金績效走勢", fund.get("tcb_url") or "")
            if not industries.empty:
                chart = _chart(zip(industries["name"], industries["weight"]), "Industry allocation")
                if chart:
                    doc.add_picture(chart, width=Inches(6.25))
                rows = [[f"#{idx}", row["name"], _fmt(row["weight"])] for idx, (_, row) in enumerate(industries.iterrows(), start=1)]
                _add_table(doc, ["圖表序號", "產業", "比重%"], rows, [1100, 6660, 1600], numeric_columns=(2,))
            else:
                doc.add_paragraph("來源未公布產業配置。")
            if not holdings.empty:
                chart = _chart(zip(holdings["name"], holdings["weight"]), "Top holdings")
                if chart:
                    doc.add_picture(chart, width=Inches(6.25))
                rows = [[f"#{idx}", row["name"], _fmt(row["weight"])] for idx, (_, row) in enumerate(holdings.iterrows(), start=1)]
                _add_table(doc, ["圖表序號", "個股", "比重%"], rows, [1100, 6660, 1600], numeric_columns=(2,))
                for _, row in holdings.iterrows():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(1)
                    _set_font(p.add_run(f"{row['name']}："), size=8)
                    _add_hyperlink(p, "Yahoo 股票技術線", stock_url(row["name"]))
            else:
                doc.add_paragraph("來源未公布主要持股。")

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("資料來源與重要聲明", level=1)
    doc.add_paragraph("資料來源包括基金公司公開資訊、合庫 MoneyDJ、Yahoo Finance 與儀表板每日更新資料。報酬、淨值及持股資料可能存在更新時間差；投資前請查閱基金公開說明書及最新月報。")
    doc.add_paragraph("近期動能加速度使用年化後的 3M 與 6M 報酬差計算，數值可能因短期波動而放大，不應單獨作為買賣依據。")

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def section_number_label(number):
    return {1: "一", 2: "二", 3: "三", 4: "四", 5: "五"}.get(number, str(number))
